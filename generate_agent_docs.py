"""
Generate docx files for each agent design doc in design/agents/.
Run: python3 generate_agent_docs.py
Output: design/agents/{AgentName}_Design.docx  (4 files)
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
TEAL   = RGBColor(0x00, 0x7A, 0x87)
SLATE  = RGBColor(0x44, 0x55, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF2, 0xF5, 0xF9)
CODE_BG = RGBColor(0xF6, 0xF8, 0xFA)
ORANGE = RGBColor(0xF4, 0x7B, 0x20)

AGENTS = [
    ("market_agent.md",       "Market_Agent_Design.docx",       "Market Agent"),
    ("research_agent.md",     "Research_Agent_Design.docx",     "Research Agent"),
    ("risk_agent.md",         "Risk_Agent_Design.docx",         "Risk Agent"),
    ("orchestrator_agent.md", "Orchestrator_Agent_Design.docx", "Orchestrator Agent"),
    ("learning_agent.md",     "Learning_Agent_Design.docx",     "Learning Agent"),
    ("news_analyst.md",       "News_Analyst_Design.docx",       "News Analyst"),
]


# ── Low-level helpers ──────────────────────────────────────────────────────────

def _shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _shade_paragraph(p, rgb: RGBColor):
    """Shade paragraph background (used for code blocks)."""
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _apply_inline(run_parent, text: str, base_size=Pt(10), base_color=None):
    """
    Add runs to `run_parent` (a paragraph) handling **bold** and `code` inline.
    """
    color = base_color or SLATE
    # Split on **bold** and `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = run_parent.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = base_size
            run.font.color.rgb = color
        elif part.startswith('`') and part.endswith('`'):
            run = run_parent.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = TEAL
        else:
            if part:
                run = run_parent.add_run(part)
                run.font.size = base_size
                run.font.color.rgb = color


# ── Document element adders ────────────────────────────────────────────────────

def add_title_block(doc, agent_name: str):
    """Cover title at top of doc."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("Trading Agent C")
    run.font.size = Pt(11)
    run.font.color.rgb = SLATE
    run.font.bold = False

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(agent_name + " — Design Doc")
    run2.font.size = Pt(18)
    run2.font.bold = True
    run2.font.color.rgb = NAVY

    # Accent line
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(16)
    run3 = p3.add_run("━" * 60)
    run3.font.color.rgb = TEAL
    run3.font.size = Pt(8)


def add_heading(doc, text: str, level: int = 2):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(14 if level == 2 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.color.rgb = NAVY if level == 2 else TEAL
    run.font.size = Pt(13 if level == 2 else 11)
    run.font.bold = True
    return p


def add_body(doc, text: str):
    """Normal paragraph with inline formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _apply_inline(p, text)
    return p


def add_meta_line(doc, text: str):
    """**Key:** value lines at the top of a doc."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0)
    # Split on first **: to separate key from value
    m = re.match(r'\*\*([^*]+)\*\*\s*(.*)', text)
    if m:
        key, value = m.group(1), m.group(2)
        r1 = p.add_run(key + ": ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
        if value:
            r2 = p.add_run(value)
            r2.font.size = Pt(10)
            r2.font.color.rgb = SLATE
    else:
        _apply_inline(p, text)
    return p


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    _apply_inline(p, text)
    return p


def add_code_block(doc, lines: list[str]):
    """Monospaced block with light background."""
    for line in lines:
        p = doc.add_paragraph()
        _shade_paragraph(p, CODE_BG)
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = SLATE
    # small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(6)


def add_table_from_rows(doc, header: list[str], rows: list[list[str]]):
    """Render a markdown table."""
    col_count = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(header):
        cell = hrow.cells[i]
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        trow = table.rows[ri + 1]
        shade = (ri % 2 == 1)
        for ci, cell_text in enumerate(row_data):
            cell = trow.cells[ci]
            if shade:
                _shade_cell(cell, LIGHT)
            p = cell.paragraphs[0]
            _apply_inline(p, cell_text.strip(), base_size=Pt(9))

    # Auto column widths — distribute evenly within page
    page_width_cm = 16.5
    col_w = page_width_cm / col_count
    _set_col_widths(table, [col_w] * col_count)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 80)
    run.font.color.rgb = LIGHT
    run.font.size = Pt(6)


# ── Markdown parser ────────────────────────────────────────────────────────────

def _parse_table_line(line: str) -> list[str]:
    """Split a markdown table row into cells."""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r'^\|?[\s\-|:]+\|?$', stripped) and '-' in stripped)


def render_markdown(doc, md_text: str):
    """
    Parse markdown line-by-line and add elements to doc.
    Handles: headings (#/##/###), meta lines (**Key:** val),
    code fences (```), tables (|...|), bullets (-), separators (---),
    and regular paragraphs.
    """
    lines = md_text.split("\n")
    i = 0
    # Skip the H1 title line — already in title block
    if lines and lines[0].startswith("# "):
        i = 1

    while i < len(lines):
        line = lines[i]

        # --- Fenced code block ---
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # --- Heading ---
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            # Second-level H1 inside doc (shouldn't happen but handle gracefully)
            add_heading(doc, line[2:].strip(), level=2)
            i += 1
            continue

        # --- Separator ---
        if re.match(r'^-{3,}$', line.strip()) or re.match(r'^={3,}$', line.strip()):
            add_separator(doc)
            i += 1
            continue

        # --- Table ---
        if "|" in line and line.strip().startswith("|"):
            # Collect all consecutive table lines
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # First line = header, second = separator, rest = data
            if len(table_lines) >= 2:
                header = _parse_table_line(table_lines[0])
                data_start = 1
                if len(table_lines) > 1 and _is_table_separator(table_lines[1]):
                    data_start = 2
                rows = [_parse_table_line(l) for l in table_lines[data_start:]]
                # Pad short rows
                rows = [r + [""] * (len(header) - len(r)) for r in rows
                        if any(c.strip() for c in r)]
                add_table_from_rows(doc, header, rows)
            continue

        # --- Bullet (- or * or numbered) ---
        m_bullet = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if m_bullet:
            indent = len(m_bullet.group(1)) // 2
            text = m_bullet.group(3)
            add_bullet(doc, text, level=indent)
            i += 1
            continue

        # --- Meta line (**Key:** value at start of doc) ---
        if re.match(r'^\*\*[A-Za-z ]+:\*\*', line.strip()):
            add_meta_line(doc, line.strip())
            i += 1
            continue

        # --- Empty line ---
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph ---
        add_body(doc, line.strip())
        i += 1


# ── Main ───────────────────────────────────────────────────────────────────────

def generate_agent_doc(src_md: Path, dest_docx: Path, agent_name: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    add_title_block(doc, agent_name)

    md_text = src_md.read_text(encoding="utf-8")
    render_markdown(doc, md_text)

    doc.save(str(dest_docx))
    print(f"Saved: {dest_docx.name}")


if __name__ == "__main__":
    base = Path(__file__).parent / "design" / "agents"
    for src_name, dest_name, label in AGENTS:
        src  = base / src_name
        dest = base / dest_name
        if not src.exists():
            print(f"SKIP (not found): {src_name}")
            continue
        generate_agent_doc(src, dest, label)

    print("\nDone — 6 agent design docs generated.")
